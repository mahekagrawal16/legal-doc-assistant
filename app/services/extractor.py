# app/services/extractor.py — Text Extraction (PDF + DOCX)
#
# WHY SEPARATE FROM OCR:
#   - Native text extraction (this file) handles digital PDFs and DOCX.
#   - OCR is needed only for scanned/image-based PDFs (handled in ocr.py).
#   - Keeping them separate makes it easy to swap either without touching the other.
#
# LIBRARIES USED:
#   - PyMuPDF (fitz): fastest and most accurate for digital PDFs; preserves structure.
#   - python-docx: official OOXML parser for .docx files.

import os


def extract_text_from_file(file_path: str) -> str:
    """
    Dispatches to the correct extractor based on file extension.
    Returns plain text string of the entire document.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _extract_from_pdf(file_path)
    elif ext == ".docx":
        return _extract_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_from_pdf(path: str) -> str:
    """
    Uses PyMuPDF for fast, structure-preserving text extraction.
    If a page yields no text (scanned page), falls back to OCR for that page.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF not installed. Run: pip install PyMuPDF")

    doc = fitz.open(path)
    pages_text = []

    for page_num, page in enumerate(doc):
        text = page.get_text("text")
        if text.strip():
            pages_text.append(text)
        else:
            # Attempt OCR fallback for image-only pages
            try:
                from app.services.ocr import ocr_pdf_page
                ocr_text = ocr_pdf_page(page)
                if ocr_text.strip():
                    pages_text.append(ocr_text)
            except Exception:
                pass  # Skip unreadable pages silently

    doc.close()
    return "\n\n".join(pages_text)


def _extract_from_docx(path: str) -> str:
    """
    Uses python-docx to extract paragraphs and table cells.
    Preserves headings and table content for better chunking.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = Document(path)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Uppercase heading-style text for the chunker to detect
            if para.style.name.startswith("Heading"):
                parts.append(f"\n{text.upper()}\n")
            else:
                parts.append(text)

    # Also extract table contents
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)